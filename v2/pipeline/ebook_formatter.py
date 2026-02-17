"""
Ebook output rendering: PDF (via WeasyPrint) and DOCX (via python-docx).

This module takes a fully populated EbookState and produces publication-ready
output files. Two renderers are provided:

  render_pdf()  -- Uses WeasyPrint (HTML/CSS to PDF) for professional layout:
    - Cover page with generated image, title, subtitle, author
    - Copyright page
    - Auto-generated table of contents with page numbers
    - Styled chapter headings and body text with full markdown support
    - Page numbers in footer, ebook title in header
    - Optional section images (centered, properly constrained)
    - Professional typography, margins, and spacing
    - Full Unicode support

  render_docx() -- Uses python-docx for Word-compatible output:
    - Cover page with title and subtitle
    - TOC field that Word auto-populates on open
    - Heading 1/2 styles for chapter/section titles
    - Consistent paragraph formatting
    - Optional inline images

Why WeasyPrint for PDF:
  - HTML/CSS gives precise, flexible layout control
  - Markdown content converts naturally to HTML
  - CSS @page rules handle headers, footers, page numbers
  - Image sizing via CSS (max-width, max-height, object-fit)
  - Full Unicode support -- no encoding workarounds needed
  - Professional typography with proper font stacks

To customize the PDF look:
  - Edit the EBOOK_CSS constant for global styling changes
  - Edit the EBOOK_HTML_TEMPLATE Jinja2 template for structural changes
"""

from __future__ import annotations

import base64
import logging
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

from v2.core.config import OUTPUT_DIR

# WeasyPrint requires system libraries (pango, cairo, gobject) that may live
# in /opt/homebrew/lib on macOS (Apple Silicon) or /usr/local/lib (Intel).
# Ensure these paths are in the dynamic library search path before WeasyPrint
# is imported for the first time.
_LIB_PATHS = ["/opt/homebrew/lib", "/usr/local/lib"]
_existing = os.environ.get("DYLD_FALLBACK_LIBRARY_PATH", "")
_needed = [p for p in _LIB_PATHS if p not in _existing and Path(p).is_dir()]
if _needed:
    os.environ["DYLD_FALLBACK_LIBRARY_PATH"] = ":".join(_needed + ([_existing] if _existing else []))

logger = logging.getLogger(__name__)


# Output directory for generated ebook files
EBOOK_OUTPUT_DIR = OUTPUT_DIR / "ebooks"
EBOOK_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _strip_markdown(text: str) -> str:
    """
    Strip common markdown formatting from text, returning plain prose.

    Used by the DOCX renderer where rich formatting isn't supported via
    markdown-to-HTML conversion. Handles bold, italic, headers, links,
    images, code blocks, and horizontal rules.
    """
    # Remove code blocks (``` ... ```)
    text = re.sub(r"```[\s\S]*?```", "", text)
    # Remove inline code (`...`)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Remove images ![alt](url)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    # Remove links [text](url) -> text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Remove headers (## text -> text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic (*** / ** / *)
    text = re.sub(r"\*{3}(.+?)\*{3}", r"\1", text)
    text = re.sub(r"\*{2}(.+?)\*{2}", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    # Remove underscore bold/italic (___ / __ / _)
    text = re.sub(r"_{3}(.+?)_{3}", r"\1", text)
    text = re.sub(r"_{2}(.+?)_{2}", r"\1", text)
    # Remove strikethrough (~~text~~)
    text = re.sub(r"~~(.+?)~~", r"\1", text)
    # Remove horizontal rules (--- or ***)
    text = re.sub(r"^[\-\*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Clean up bullet list markers at start of lines
    text = re.sub(r"^\s*[\-\*\+]\s+", "", text, flags=re.MULTILINE)
    # Clean up numbered list markers at start of lines
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    return text.strip()


def _image_to_data_uri(image_path: Path) -> Optional[str]:
    """
    Convert an image file to a base64 data URI for embedding in HTML.

    Returns None if the file doesn't exist or can't be read.
    """
    if not image_path or not image_path.exists():
        return None
    try:
        suffix = image_path.suffix.lower()
        mime_types = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".webp": "image/webp",
        }
        mime = mime_types.get(suffix, "image/png")
        data = image_path.read_bytes()
        b64 = base64.b64encode(data).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception as e:
        logger.warning(f"Failed to encode image {image_path}: {e}")
        return None


# ═══════════════════════════════════════════════════════════════════════════
# PDF RENDERER (WeasyPrint)
# ═══════════════════════════════════════════════════════════════════════════

EBOOK_CSS = """
/* ── Page setup ── */
@page {
    size: A4;
    margin: 25mm 25mm 30mm 25mm;

    @bottom-center {
        content: counter(page);
        font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
        font-size: 9pt;
        color: #999;
    }

    @top-left {
        content: string(ebook-title);
        font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
        font-size: 8pt;
        font-style: italic;
        color: #999;
    }

    @top-right {
        content: "";
        border-bottom: 0.5pt solid #ccc;
    }

    @top-left {
        border-bottom: 0.5pt solid #ccc;
    }

    @top-center {
        border-bottom: 0.5pt solid #ccc;
    }
}

/* No header/footer on cover, copyright, or TOC pages */
@page cover {
    @top-left { content: none; border: none; }
    @top-center { border: none; }
    @top-right { content: none; border: none; }
    @bottom-center { content: none; }
}

@page copyright {
    @top-left { content: none; border: none; }
    @top-center { border: none; }
    @top-right { content: none; border: none; }
    @bottom-center { content: none; }
}

@page toc {
    @top-left { content: none; border: none; }
    @top-center { border: none; }
    @top-right { content: none; border: none; }
}

/* ── Base typography ── */
html {
    font-size: 11pt;
}

body {
    font-family: Georgia, "Times New Roman", Times, serif;
    color: #282828;
    line-height: 1.65;
    word-wrap: break-word;
    overflow-wrap: break-word;
}

/* ── Cover page ── */
.cover-page {
    page: cover;
    page-break-after: always;
    text-align: center;
    display: flex;
    flex-direction: column;
    justify-content: flex-start;
    align-items: center;
    min-height: 100%;
}

.cover-image-wrapper {
    width: 100%;
    text-align: center;
    margin-top: 10mm;
    margin-bottom: 15mm;
}

.cover-image {
    max-width: 100%;
    max-height: 150mm;
    object-fit: contain;
    display: block;
    margin: 0 auto;
}

.cover-title {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 28pt;
    font-weight: bold;
    color: #1a1a1a;
    margin: 0 0 8mm 0;
    line-height: 1.2;
    string-set: ebook-title content();
}

.cover-subtitle {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 15pt;
    font-style: italic;
    color: #555;
    margin: 0 0 10mm 0;
    line-height: 1.3;
}

.cover-author {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 13pt;
    color: #777;
    margin: 0;
}

/* ── Copyright page ── */
.copyright-page {
    page: copyright;
    page-break-after: always;
    text-align: center;
    padding-top: 45%;
}

.copyright-page p {
    font-size: 9pt;
    color: #888;
    line-height: 1.8;
    margin: 0;
}

/* ── Table of Contents ── */
.toc-page {
    page: toc;
    page-break-after: always;
}

.toc-title {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 22pt;
    font-weight: bold;
    color: #1a1a1a;
    text-align: center;
    margin-bottom: 12mm;
    padding-bottom: 4mm;
    border-bottom: 1px solid #ddd;
}

.toc-list {
    list-style: none;
    padding: 0;
    margin: 0;
}

.toc-chapter {
    margin-bottom: 2mm;
}

.toc-chapter a {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 11pt;
    font-weight: bold;
    color: #282828;
    text-decoration: none;
    display: block;
    padding: 2mm 0;
}

.toc-chapter a::after {
    content: target-counter(attr(href), page);
    float: right;
    font-weight: normal;
    color: #666;
}

.toc-section {
    margin-bottom: 1mm;
    padding-left: 15mm;
}

.toc-section a {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 10pt;
    font-weight: normal;
    color: #555;
    text-decoration: none;
    display: block;
    padding: 1mm 0;
}

.toc-section a::after {
    content: target-counter(attr(href), page);
    float: right;
    color: #888;
}

/* ── Chapter pages ── */
.chapter {
    page-break-before: always;
}

.chapter-number {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 14pt;
    font-weight: bold;
    color: #999;
    text-transform: uppercase;
    letter-spacing: 2pt;
    margin: 0 0 2mm 0;
}

.chapter-title {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 24pt;
    font-weight: bold;
    color: #1a1a1a;
    margin: 0 0 8mm 0;
    line-height: 1.2;
    padding-bottom: 4mm;
    border-bottom: 2px solid #ddd;
}

.chapter-image-wrapper {
    text-align: center;
    margin: 8mm 0;
    page-break-inside: avoid;
}

.chapter-image {
    max-width: 65%;
    max-height: 80mm;
    object-fit: contain;
    display: block;
    margin: 0 auto;
    border-radius: 2px;
}

/* ── Section headings ── */
.section-title {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 15pt;
    font-weight: bold;
    color: #333;
    margin: 8mm 0 4mm 0;
    page-break-after: avoid;
}

/* ── Body text ── */
.section-body {
    margin-bottom: 4mm;
}

.section-body p {
    margin: 0 0 3mm 0;
    text-align: justify;
    hyphens: auto;
}

.section-body strong {
    font-weight: bold;
}

.section-body em {
    font-style: italic;
}

.section-body ul, .section-body ol {
    margin: 2mm 0 4mm 8mm;
    padding-left: 4mm;
}

.section-body li {
    margin-bottom: 2mm;
}

.section-body blockquote {
    border-left: 3px solid #ccc;
    margin: 4mm 0 4mm 4mm;
    padding: 2mm 4mm 2mm 8mm;
    color: #555;
    font-style: italic;
}

/* ── Introduction & Conclusion ── */
.intro-page, .conclusion-page {
    page-break-before: always;
}

.intro-title, .conclusion-title {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 22pt;
    font-weight: bold;
    color: #1a1a1a;
    margin: 0 0 8mm 0;
    padding-bottom: 4mm;
    border-bottom: 2px solid #ddd;
}

.intro-body, .conclusion-body {
    margin-bottom: 4mm;
}

.intro-body p, .conclusion-body p {
    margin: 0 0 3mm 0;
    text-align: justify;
    hyphens: auto;
}

/* ── Prevent orphans/widows ── */
p {
    orphans: 3;
    widows: 3;
}

h1, h2, h3 {
    page-break-after: avoid;
}

img {
    page-break-inside: avoid;
}
"""

EBOOK_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <style>{{ css }}</style>
</head>
<body>

    <!-- Cover Page -->
    <div class="cover-page">
        {% if cover_image_uri %}
        <div class="cover-image-wrapper">
            <img class="cover-image" src="{{ cover_image_uri }}" alt="Cover">
        </div>
        {% endif %}
        <h1 class="cover-title">{{ ebook_title }}</h1>
        {% if ebook_subtitle %}
        <p class="cover-subtitle">{{ ebook_subtitle }}</p>
        {% endif %}
        <p class="cover-author">by {{ author_name }}</p>
    </div>

    <!-- Copyright Page -->
    <div class="copyright-page">
        <p>&copy; {{ year }} {{ author_name }}. All rights reserved.</p>
        <p>&nbsp;</p>
        <p>No part of this publication may be reproduced, distributed,<br>
        or transmitted in any form without prior written permission.</p>
        <p>&nbsp;</p>
        <p>Generated by MoceanAI &mdash; {{ generated_date }}.</p>
    </div>

    <!-- Table of Contents -->
    <div class="toc-page">
        <h1 class="toc-title">Table of Contents</h1>
        <ul class="toc-list">
            {% if has_introduction %}
            <li class="toc-chapter"><a href="#introduction">Introduction</a></li>
            {% endif %}
            {% for ch in chapters %}
            <li class="toc-chapter">
                <a href="#chapter-{{ ch.index }}">Chapter {{ ch.index }}: {{ ch.title }}</a>
            </li>
            {% for sec in ch.sections %}
            <li class="toc-section">
                <a href="#section-{{ ch.index }}-{{ loop.index }}">{{ sec.title }}</a>
            </li>
            {% endfor %}
            {% endfor %}
            {% if has_conclusion %}
            <li class="toc-chapter"><a href="#conclusion">Conclusion</a></li>
            {% endif %}
        </ul>
    </div>

    <!-- Introduction -->
    {% if introduction_html %}
    <div class="intro-page" id="introduction">
        <h1 class="intro-title">Introduction</h1>
        <div class="intro-body">{{ introduction_html }}</div>
    </div>
    {% endif %}

    <!-- Chapters -->
    {% for ch in chapters %}
    <div class="chapter" id="chapter-{{ ch.index }}">
        <p class="chapter-number">Chapter {{ ch.index }}</p>
        <h1 class="chapter-title">{{ ch.title }}</h1>

        {% if ch.image_uri %}
        <div class="chapter-image-wrapper">
            <img class="chapter-image" src="{{ ch.image_uri }}" alt="Chapter {{ ch.index }} illustration">
        </div>
        {% endif %}

        {% for sec in ch.sections %}
        <h2 class="section-title" id="section-{{ ch.index }}-{{ loop.index }}">{{ sec.title }}</h2>
        <div class="section-body">{{ sec.html }}</div>
        {% endfor %}
    </div>
    {% endfor %}

    <!-- Conclusion -->
    {% if conclusion_html %}
    <div class="conclusion-page" id="conclusion">
        <h1 class="conclusion-title">Conclusion</h1>
        <div class="conclusion-body">{{ conclusion_html }}</div>
    </div>
    {% endif %}

</body>
</html>
"""


def render_pdf(state) -> Path:
    """
    Render the ebook as a professionally styled PDF using WeasyPrint.

    Converts markdown content to HTML, renders it with a Jinja2 template
    and professional CSS, then generates the PDF via WeasyPrint.

    Args:
        state: Fully populated EbookState.

    Returns:
        Path to the generated .pdf file.
    """
    import markdown as md
    from jinja2 import Template
    from weasyprint import HTML

    config = state.config
    outline = state.outline
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c if c.isalnum() or c in " _-" else "" for c in config.title)[:50]
    output_path = EBOOK_OUTPUT_DIR / f"{safe_title}_{timestamp}.pdf"

    md_extensions = ["extra", "smarty", "sane_lists"]

    def _md_to_html(text: str) -> str:
        """Convert markdown text to HTML, handling paragraph breaks."""
        if not text:
            return ""
        return md.markdown(text, extensions=md_extensions)

    # Build cover image data URI
    cover_image_uri = None
    if state.cover_image_path:
        cover_image_uri = _image_to_data_uri(state.cover_image_path)

    # Build chapter data for the template
    chapters_data = []
    for ch_idx, chapter in enumerate(state.chapters):
        sections = chapter.edited_sections or chapter.raw_sections
        if not sections:
            continue

        # Chapter image
        ch_image_uri = None
        if chapter.section_image_paths:
            for img_path in chapter.section_image_paths:
                uri = _image_to_data_uri(img_path)
                if uri:
                    ch_image_uri = uri
                    break

        sections_data = []
        for section in sections:
            sections_data.append({
                "title": section.section_title,
                "html": _md_to_html(section.section_text),
            })

        chapters_data.append({
            "index": ch_idx + 1,
            "title": chapter.outline.chapter_title,
            "image_uri": ch_image_uri,
            "sections": sections_data,
        })

    # Render the HTML template
    template = Template(EBOOK_HTML_TEMPLATE)
    ebook_title = outline.ebook_title if outline else config.title
    ebook_subtitle = (outline.ebook_subtitle if outline else config.subtitle) or ""

    html_content = template.render(
        css=EBOOK_CSS,
        cover_image_uri=cover_image_uri,
        ebook_title=ebook_title,
        ebook_subtitle=ebook_subtitle,
        author_name=config.author_name,
        year=datetime.now().year,
        generated_date=datetime.now().strftime("%B %d, %Y"),
        has_introduction=bool(state.introduction_text),
        introduction_html=_md_to_html(state.introduction_text or ""),
        chapters=chapters_data,
        has_conclusion=bool(state.conclusion_text),
        conclusion_html=_md_to_html(state.conclusion_text or ""),
    )

    # Generate the PDF
    html_doc = HTML(string=html_content)
    html_doc.write_pdf(str(output_path))

    logger.info(f"PDF rendered: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════
# DOCX RENDERER (python-docx)
# ═══════════════════════════════════════════════════════════════════════════

def render_docx(state) -> Path:
    """
    Render the ebook as a DOCX file using python-docx.

    The document uses Word heading styles (Heading 1, Heading 2) so that
    Word can auto-generate a table of contents. A TOC field is inserted
    that will be populated when the user opens the file in Word.

    All text is run through _strip_markdown() to remove any markdown
    formatting that the LLM may have included.

    Args:
        state: Fully populated EbookState.

    Returns:
        Path to the generated .docx file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    config = state.config
    outline = state.outline
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = "".join(c if c.isalnum() or c in " _-" else "" for c in config.title)[:50]
    output_path = EBOOK_OUTPUT_DIR / f"{safe_title}_{timestamp}.docx"

    doc = Document()

    # ── Cover page ──
    ebook_title = outline.ebook_title if outline else config.title
    subtitle_text = (outline.ebook_subtitle if outline else config.subtitle) or ""

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(ebook_title)
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(30, 30, 30)

    # Subtitle
    if subtitle_text:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle_text)
        sub_run.font.size = Pt(16)
        sub_run.italic = True
        sub_run.font.color.rgb = RGBColor(80, 80, 80)

    # Author
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"by {config.author_name}")
    author_run.font.size = Pt(14)
    author_run.font.color.rgb = RGBColor(100, 100, 100)

    # Cover image
    if state.cover_image_path and state.cover_image_path.exists():
        try:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.add_run().add_picture(str(state.cover_image_path), width=Inches(4))
        except Exception as e:
            logger.warning(f"DOCX cover image failed: {e}")

    doc.add_page_break()

    # ── Copyright page ──
    year = datetime.now().year
    copyright_text = (
        f"\u00a9 {year} {config.author_name}. All rights reserved.\n\n"
        f"No part of this publication may be reproduced, distributed, "
        f"or transmitted in any form without prior written permission.\n\n"
        f"Generated by MoceanAI \u2014 {datetime.now().strftime('%B %d, %Y')}."
    )
    copyright_para = doc.add_paragraph(copyright_text)
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in copyright_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ── Table of Contents field ──
    toc_heading = doc.add_heading("Table of Contents", level=1)  # noqa: F841
    _insert_toc_field(doc)
    doc.add_page_break()

    # ── Introduction ──
    if state.introduction_text:
        doc.add_heading("Introduction", level=1)
        cleaned = _strip_markdown(state.introduction_text)
        for paragraph in cleaned.split("\n\n"):
            clean = " ".join(paragraph.split())
            if clean:
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(6)
        doc.add_page_break()

    # ── Chapters ──
    for ch_idx, chapter in enumerate(state.chapters):
        sections = chapter.edited_sections or chapter.raw_sections
        if not sections:
            continue

        chapter_title = chapter.outline.chapter_title
        doc.add_heading(f"Chapter {ch_idx + 1}: {chapter_title}", level=1)

        # Chapter image
        if chapter.section_image_paths:
            for img_path in chapter.section_image_paths:
                if img_path.exists():
                    try:
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.add_run().add_picture(str(img_path), width=Inches(4))
                    except Exception as e:
                        logger.warning(f"DOCX chapter image failed: {e}")

        for section in sections:
            doc.add_heading(_strip_markdown(section.section_title), level=2)
            cleaned = _strip_markdown(section.section_text)
            for paragraph in cleaned.split("\n\n"):
                clean = " ".join(paragraph.split())
                if clean:
                    p = doc.add_paragraph(clean)
                    p.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    # ── Conclusion ──
    if state.conclusion_text:
        doc.add_heading("Conclusion", level=1)
        cleaned = _strip_markdown(state.conclusion_text)
        for paragraph in cleaned.split("\n\n"):
            clean = " ".join(paragraph.split())
            if clean:
                p = doc.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(6)

    # ── Save ──
    doc.save(str(output_path))
    logger.info(f"DOCX rendered: {output_path}")
    return output_path


def _insert_toc_field(doc) -> None:
    """
    Insert a Word TOC field code into the document.

    This creates a field that Word will populate with actual page numbers
    when the user opens the document. The user will see a prompt to
    "Update Table" which fills in the entries.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and select 'Update Field' to generate Table of Contents"
    run._r.append(fld_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
